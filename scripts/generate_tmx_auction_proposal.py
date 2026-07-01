from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def bullet_paragraph(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def main():
    out_path = r"c:\Projects\WRRB-TRAINING-SYSTEM-v1\TMX_Auction_Project_Proposal.docx"

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    # Title
    p = doc.add_paragraph("Project Proposal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(20)

    p2 = doc.add_paragraph("TMX Auction Data Integration (Mode B: TMX → WRRB)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(16)

    p3 = doc.add_paragraph("WRRB Training System v1 (Laravel)")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_paragraph("Prepared for: WRRB / TMX Integration Stakeholders")
    doc.add_paragraph("Prepared by: Development Team")
    doc.add_paragraph(f"Date: {date.today().isoformat()}")

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This proposal covers the implementation of TMX Auction data integration for the WRRB Training System v1. "
        "The system supports Mode B (Manual Pull Trigger / On-demand re-delivery), where WRRB calls TMX endpoints "
        "to force immediate re-export of auction updates. The integration includes an OAuth 2.0 client_credentials "
        "flow for token generation, a super_admin-only interactive dashboard to view fetched auction objects, "
        "persistence of fetched snapshots into a normalized relational database schema for auditing, and a CSV "
        "export for operational reporting and reconciliation."
    )

    # 2. Background
    doc.add_heading("2. Background and Problem Statement", level=1)
    doc.add_paragraph(
        "WRRB requires a reliable and auditable way to obtain auction lifecycle information from TMX so that WRRB can "
        "track auction outcomes for receipts. Because polling and re-delivery are required to recover from downtime "
        "and to support reconciliation, the integration must provide clear operational visibility and must persist "
        "data snapshots."
    )

    # 3. Objectives
    doc.add_heading("3. Objectives", level=1)
    bullet_paragraph(
        doc,
        "Implement Mode B TMX → WRRB integration using TMX OAuth 2.0 token endpoint (client_credentials) and "
        "Bearer token authentication."
    )
    bullet_paragraph(
        doc,
        "Provide a super_admin-only dashboard to fetch auction data using filters (from_date, to_date, receipt_no)."
    )
    bullet_paragraph(
        doc,
        "Persist all fetched payloads into a normalized database schema (deliveries, auctions, lots, receipts, and history)."
    )
    bullet_paragraph(
        doc,
        "Display the latest known state per auction (including winner, consent, release order, and receipt status) after subsequent fetches."
    )
    bullet_paragraph(
        doc,
        "Enable CSV downloads of the latest saved data for reporting and reconciliation."
    )

    # 4. Scope
    doc.add_heading("4. Scope", level=1)
    doc.add_paragraph("In scope:")
    bullet_paragraph(doc, "TMX Mode B OAuth token request to /integration/oauth/token using client_credentials.")
    bullet_paragraph(doc, "TMX Mode B endpoints: /integration/wrrb/v1/export-pending and /integration/wrrb/v1/export-by-receipt/{receipt_no}.")
    bullet_paragraph(doc, "Dashboard UI: filter form and expandable table view for auctions → lots → receipts.")
    bullet_paragraph(doc, "Database persistence: migrations, models, and service logic to store fetched snapshots.")
    bullet_paragraph(doc, "CSV export endpoint to download latest saved data per auction.")
    doc.add_paragraph("Out of scope (for this phase):")
    bullet_paragraph(doc, "Mode A webhook receiving (TMX → WRRB push).")
    bullet_paragraph(doc, "Automated nightly reconciliation beyond manual pull triggers.")

    # 5. Functional Design
    doc.add_heading("5. Functional Design", level=1)
    doc.add_heading("5.1 Authentication (OAuth 2.0)", level=2)
    doc.add_paragraph(
        "The system obtains a Bearer token from TMX using the OAuth 2.0 client_credentials grant. The token is then "
        "used as Authorization: Bearer <token> on all Mode B export calls."
    )

    doc.add_heading("5.2 Mode B Endpoints", level=2)
    bullet_paragraph(
        doc,
        "Export Pending: POST /integration/wrrb/v1/export-pending (optional JSON body: from_date, to_date, receipt_no)."
    )
    bullet_paragraph(
        doc,
        "Export by Receipt: POST /integration/wrrb/v1/export-by-receipt/{receipt_no} (Bearer token; no body required)."
    )

    doc.add_heading("5.3 Dashboard", level=2)
    doc.add_paragraph(
        "The dashboard is accessible only to super_admin. Users can filter and fetch auction data. The UI uses an "
        "interactive expandable table to show auction-level fields and, on expansion, lot-level and receipt-level "
        "details including winner, consent, release order, and lifecycle history."
    )

    # 6. Data Persistence Design
    doc.add_heading("6. Data Persistence Design", level=1)
    doc.add_paragraph(
        "The TMX payload is stored in separate related tables for easier management:"
    )
    bullet_paragraph(doc, "tmx_deliveries: top-level snapshots keyed by request_id (TMX request_id / X-Request-ID).")
    bullet_paragraph(doc, "tmx_auctions: auctions under a delivery.")
    bullet_paragraph(doc, "tmx_auction_lots: lots under an auction, including commodity, winner, consent, and release order fields.")
    bullet_paragraph(doc, "tmx_auction_receipts: receipts under a lot (one lot can have multiple receipts).")
    bullet_paragraph(doc, "tmx_auction_lot_history: lifecycle history entries under a lot.")

    doc.add_heading("7. Latest State Strategy", level=1)
    doc.add_paragraph(
        "TMX may deliver updated payloads for the same auction over time. To ensure the dashboard stays useful, "
        "the system computes the latest state per auction from the database (ordered by tmx_updated_at) so fields "
        "like winner and release order reflect the newest information."
    )

    # 8. CSV Export
    doc.add_heading("8. CSV Export", level=1)
    doc.add_paragraph(
        "A super_admin-only button downloads a CSV file containing the latest saved auction state per auction, "
        "expanded down to receipts and key operational fields. This supports offline review and reconciliation."
    )

    # 9. Security
    doc.add_heading("9. Security and Access Control", level=1)
    bullet_paragraph(doc, "Dashboard and export endpoints are protected by super_admin middleware.")
    bullet_paragraph(doc, "Secrets are never written in full to logs; debug logging focuses on presence/length and safe metadata.")
    bullet_paragraph(doc, "SSL verification is configurable for non-production environments; it must be enabled in production.")

    # 10. Testing Plan
    doc.add_heading("10. Testing and Validation Plan", level=1)
    bullet_paragraph(doc, "Manual fetch tests: export-pending without filters; with from/to date; and export-by-receipt for a known receipt_no.")
    bullet_paragraph(doc, "Authentication validation: verify 404/401 error handling and correct token request format per TMX spec.")
    bullet_paragraph(doc, "Persistence validation: confirm deliveries/auctions/lots/receipts/history are stored correctly and replaced/upserted per request_id.")
    bullet_paragraph(doc, "Dashboard validation: confirm the latest winner/consent/release-order appears after subsequent fetches.")
    bullet_paragraph(doc, "Export validation: ensure CSV rows match the displayed latest state and open correctly in spreadsheet software.")

    # 11. Timeline
    doc.add_heading("11. Implementation Timeline (Suggested)", level=1)
    bullet_paragraph(doc, "Day 1-2: finalize endpoint details, configure environment variables, and implement token + Mode B requests with error handling.")
    bullet_paragraph(doc, "Day 3-4: build dashboard UI (filters + interactive table) and ensure correct rendering of nested auction payloads.")
    bullet_paragraph(doc, "Day 5-6: create database schema (migrations/models) and persistence service; implement saving after successful fetch.")
    bullet_paragraph(doc, "Day 7: implement latest state per auction logic and CSV export endpoint.")
    bullet_paragraph(doc, "Day 8-10: end-to-end manual testing, edge case handling, stabilization, and documentation.")

    # 12. Risks
    doc.add_heading("12. Risks and Mitigations", level=1)
    bullet_paragraph(doc, "Risk: Incorrect endpoint paths or base URL causes 404. Mitigation: configurable paths and error messages with failing URLs.")
    bullet_paragraph(doc, "Risk: OAuth 401 unauthorized due to credential or request format. Mitigation: strict JSON payload per TMX spec and safe debug logs.")
    bullet_paragraph(doc, "Risk: Payload structure variations. Mitigation: tolerant token extraction and integration testing with TMX samples.")
    bullet_paragraph(doc, "Risk: Data duplication or inconsistency across fetches. Mitigation: request_id-based delivery upsert and latest-per-auction display.")

    # Final
    doc.add_heading("13. Conclusion", level=1)
    doc.add_paragraph(
        "This implementation will provide a reliable TMX Mode B integration with operational visibility via an interactive "
        "dashboard, persistent storage for auditing and reconciliation, and an export mechanism for operational use."
    )

    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()

